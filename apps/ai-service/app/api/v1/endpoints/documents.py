import base64
import datetime
import io
import json
import os
import re
import sys
import time
import threading
import traceback
import uuid
from functools import lru_cache
from typing import Any, Dict, List, Optional, Union

import docx
import fitz
import pandas as pd
from app.utils.document_profile import (
    empty_profile,
    extract_profile,
    render_chunk_prefix,
    render_summary_chunk_text,
)
from app.utils.logger import log_to_backend
from app.utils.pdf_ocr import extract_pdf_page_texts, run_ocr_on_pil
from app.utils.processing_abort import (
    REASON_CANCELLED,
    REASON_TIMEOUT,
    ProcessingAborted,
)
from cryptography.hazmat.backends import default_backend
from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.primitives.ciphers.aead import AESGCM
from cryptography.hazmat.primitives.kdf.hkdf import HKDF
from fastapi import APIRouter, Depends, File, Form, Header, HTTPException, UploadFile
from fastapi.responses import StreamingResponse
from fastapi.security import HTTPAuthorizationCredentials, HTTPBearer
from PIL import Image
from pydantic import BaseModel
from sentence_transformers import SentenceTransformer

# Add the root directory to Python path
backend_dir = os.path.dirname(
    os.path.dirname(
        os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    )
)
if backend_dir not in sys.path:
    sys.path.append(backend_dir)

# Now we can import from absolute paths
from app.core.config import settings

# OCR Configuration (see app.core.config / pdf_ocr — single source: settings)
OCR_ENABLED = settings.OCR_ENABLED
OCR_LANGUAGE = settings.OCR_LANGUAGE
OCR_DPI = settings.OCR_DPI

# Set Tesseract path if needed (uncomment and set path if tesseract is not in PATH)
# import pytesseract
# pytesseract.pytesseract.tesseract_cmd = '/opt/homebrew/bin/tesseract'

# Initialize the embedding model with the same model as query service
try:
    embedding_model = SentenceTransformer("BAAI/bge-small-en-v1.5", device="cpu")
    print("✅ Embedding model initialized successfully")
except Exception as e:
    print(f"Warning: Could not initialize SentenceTransformer: {str(e)}")
    embedding_model = None

# Processing status tracking
processing_status = {}  # Global dict to track processing status
processing_lock = threading.Lock()  # Thread lock for status updates
# Cooperative cancel: DELETE in Nest calls /cancel/{aiDocumentId}; worker checks between pages/chunks
processing_cancel_requests: set[str] = set()


def request_cancel_processing(document_id: str) -> None:
    with processing_lock:
        processing_cancel_requests.add(document_id)


def clear_cancel_processing(document_id: str) -> None:
    with processing_lock:
        processing_cancel_requests.discard(document_id)


def is_processing_cancelled(document_id: str) -> bool:
    with processing_lock:
        return document_id in processing_cancel_requests


def make_abort_guard(document_id: str, deadline: float):
    """Raise ProcessingAborted on user cancel or wall-clock timeout."""

    def guard() -> None:
        if is_processing_cancelled(document_id):
            raise ProcessingAborted(REASON_CANCELLED)
        if time.time() > deadline:
            raise ProcessingAborted(REASON_TIMEOUT)

    return guard

# Initialize Qdrant client
try:
    from qdrant_client import QdrantClient
    from qdrant_client.models import Distance, PointStruct, VectorParams

    # HTTP timeout for ALL Qdrant operations including upsert. Default httpx
    # timeout (5s write) is too short for bulk upsert over remote networks —
    # we have seen 500+ point upserts time out on Qdrant Cloud. Give it room.
    _qdrant_timeout = int(getattr(settings, "QDRANT_CLIENT_TIMEOUT_SEC", 120) or 120)

    # Initialize Qdrant client with optional authentication
    if settings.QDRANT_API_KEY:
        qdrant_client = QdrantClient(
            url=settings.QDRANT_URL,
            api_key=settings.QDRANT_API_KEY,
            timeout=_qdrant_timeout,
            check_compatibility=False,  # Skip version compatibility check
        )
    else:
        qdrant_client = QdrantClient(
            settings.QDRANT_URL,
            timeout=_qdrant_timeout,
            check_compatibility=False,  # Skip version compatibility check
        )

    qdrant_collection = settings.QDRANT_COLLECTION
    print(
        f"✅ Qdrant client initialized successfully for collection: "
        f"{qdrant_collection} (timeout={_qdrant_timeout}s)"
    )

except Exception as e:
    print(f"Warning: Could not initialize Qdrant client: {str(e)}")
    qdrant_client = None


def update_processing_status(
    document_id: str,
    status: str,
    details: str = None,
    error: str = None,
    progress: int = None,
):
    """Update processing status for a document."""
    with processing_lock:
        processing_status[document_id] = {
            "status": status,
            "details": details,
            "error": error,
            "progress": progress,  # 0-100 percentage
            "timestamp": datetime.datetime.now().isoformat(),
        }


def get_processing_status(document_id: str) -> Dict[str, Any]:
    """Get processing status for a document."""
    with processing_lock:
        return processing_status.get(document_id, {"status": "NOT_FOUND"})


def ensure_qdrant_collection_exists():
    """
    Ensure Qdrant collection exists; create it if missing (Qdrant Cloud compatible).
    Returns {"ok": bool, "detail": str | None} — detail is safe to return in HTTP responses (no secrets).
    """
    coll = qdrant_collection
    ctx = f"Qdrant collection={coll!r}"

    if not qdrant_client:
        msg = (
            f"{ctx}: client not initialized. Set QDRANT_URL (and QDRANT_API_KEY for Qdrant Cloud) "
            "in the AI service environment."
        )
        print(f"❌ {msg}")
        return {"ok": False, "detail": msg}

    def _collection_exists(name: str) -> bool:
        # Use get_collections() only — collection_exists() hits /collections/{name}/exists which
        # can return 404 on some Qdrant Cloud versions even when the cluster URL is valid.
        try:
            names = [c.name for c in qdrant_client.get_collections().collections]
            return name in names
        except Exception as ex:
            print(f"❌ {ctx}: failed to list collections: {ex}")
            raise

    try:
        exists = _collection_exists(coll)
        if exists:
            print(f"✅ Collection '{coll}' already exists")
            return {"ok": True, "detail": None}

        print(f"❌ Collection '{coll}' not found; creating…")
        # BAAI/bge-small-en-v1.5 → 384 dimensions, cosine distance
        qdrant_client.create_collection(
            collection_name=coll,
            vectors_config=VectorParams(size=384, distance=Distance.COSINE),
        )
        print(f"✅ Created collection '{coll}'")

        for field_name in ("user_id", "document_id"):
            try:
                qdrant_client.create_payload_index(
                    collection_name=coll,
                    field_name=field_name,
                    field_schema="keyword",
                )
                print(f"✅ Payload index on {field_name!r} for {coll!r}")
            except Exception as idx_ex:
                print(f"⚠️  Could not create payload index {field_name!r}: {idx_ex}")

        print(f"🎉 Collection '{coll}' setup completed successfully!")
        return {"ok": True, "detail": None}

    except Exception as e:
        tb = traceback.format_exc()
        print(f"❌ {ctx}: {type(e).__name__}: {e}\n{tb}")
        err_s = f"{e}"
        base = (
            f"{ctx}: {type(e).__name__}: {e}. "
            "Ensure QDRANT_COLLECTION matches the Nest backend. "
        )
        if (
            "404" in err_s
            or "404 page not found" in err_s.lower()
            or ("not found" in err_s.lower() and "Unexpected Response" in err_s)
        ):
            base += (
                "HTTP 404 usually means QDRANT_URL points at a deleted or wrong cluster — open "
                "Qdrant Cloud → Clusters → your cluster → Connect, copy the HTTPS endpoint "
                "(port 6333), and set QDRANT_URL + QDRANT_API_KEY identically in apps/ai-service/.env "
                "and apps/backend/.env. Then restart the AI service."
            )
        else:
            base += (
                "Verify QDRANT_URL, QDRANT_API_KEY, and network access to Qdrant Cloud."
            )
        return {"ok": False, "detail": base}


# Cache for embeddings using lru_cache
@lru_cache(maxsize=1000)
def get_cached_embedding(text: str) -> List[float]:
    """Get cached embedding or compute new one."""
    return embedding_model.encode(text).tolist()


def optimize_chunk_size(text: str, target_size: int = 800) -> List[str]:
    """Optimize text into chunks of approximately target_size characters."""
    print(f"🔧 optimize_chunk_size called:")
    print(f"  Input text length: {len(text)}")
    print(f"  Target size: {target_size}")
    print(f"  Text is empty after strip: {not text.strip()}")

    if not text.strip():
        print(f"  ⚠️ Returning empty chunks - text is empty after strip")
        return []

    # Split by sentences first
    sentences = re.split(r"(?<=[.!?])\s+", text.strip())
    print(f"  Split into {len(sentences)} sentences")

    chunks = []
    current_chunk = ""

    for sentence in sentences:
        if len(current_chunk) + len(sentence) <= target_size:
            current_chunk += sentence + " "
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            current_chunk = sentence + " "

    if current_chunk:
        chunks.append(current_chunk.strip())

    return chunks


def parse_explicit_page_markers(text: str) -> List[Dict[str, Union[int, str]]]:
    """Parse OCR page markers like [[PAGE_1]] from plain text payloads."""
    pattern = re.compile(
        r"\[\[PAGE_(\d{1,5})\]\]\s*([\s\S]*?)(?=\s*\[\[PAGE_\d{1,5}\]\]|\s*$)"
    )
    pages: List[Dict[str, Union[int, str]]] = []
    for match in pattern.finditer(text or ""):
        page_num_raw = match.group(1)
        page_text = (match.group(2) or "").strip()
        if not page_num_raw or not page_text:
            continue
        try:
            page_num = int(page_num_raw)
        except Exception:
            continue
        if page_num <= 0:
            continue
        pages.append({"page_number": page_num, "text": page_text})

    pages.sort(key=lambda p: int(p["page_number"]))  # type: ignore[index]
    return pages


def optimize_chunk_size_with_pages(
    text: str,
    file_extension: str,
    file_content: bytes,
    document_id: Optional[str] = None,
    pdf_page_texts: Optional[List[str]] = None,
) -> List[Dict[str, Any]]:
    """Optimize text into chunks with page tracking for better search results.

    For PDFs, pass ``pdf_page_texts`` (cleaned, one string per page) from
    :func:`extract_pdf_page_texts` so chunking matches ingestion OCR exactly.
    """
    print(f"🔧 optimize_chunk_size_with_pages called:")
    print(f"  Input text length: {len(text)}")
    print(f"  Input text preview: {text[:100]}...")
    print(f"  File extension: {file_extension}")
    print(f"  Text is empty after strip: {not text.strip()}")

    chunks: List[Dict[str, Any]] = []

    if file_extension.lower() != ".pdf" and not (text or "").strip():
        print("  ⚠️ Returning empty chunks - text is empty after strip")
        return []

    if file_extension.lower() == ".pdf":
        print("  📄 Processing as PDF (unified OCR path)")
        try:
            if pdf_page_texts is None:
                pdf_page_texts = [
                    clean_text_content(p) for p in extract_pdf_page_texts(file_content)
                ]
            else:
                pdf_page_texts = [clean_text_content(p) for p in pdf_page_texts]

            pos_text = "\n\n".join(pdf_page_texts)
            if not pos_text.strip():
                print("  ⚠️ No PDF text after extraction")
                return []

            current_pos = 0
            num_pages = len(pdf_page_texts)

            for page_num in range(num_pages):
                if document_id and num_pages > 0:
                    pct = 80 + int(((page_num + 1) / num_pages) * 9)
                    update_processing_status(
                        document_id,
                        "PROCESSING",
                        f"Chunking PDF page {page_num + 1}/{num_pages}…",
                        None,
                        min(pct, 89),
                    )

                page_text = pdf_page_texts[page_num]
                if not page_text or not page_text.strip():
                    print(
                        f"  📄 Page {page_num + 1}: Skipping (no text from any source)"
                    )
                    continue

                print(
                    f"  📄 Page {page_num + 1}: {len(page_text)} characters (unified extract)"
                )

                page_start = pos_text.find(page_text, current_pos)
                if page_start == -1:
                    page_start = current_pos
                    print(f"  📄 Page {page_num + 1}: Using fallback position")

                page_end = page_start + len(page_text)
                current_pos = page_end
                print(f"  📄 Page {page_num + 1}: Position {page_start}-{page_end}")

                print(f"  📄 Page {page_num + 1}: Calling optimize_chunk_size")
                page_chunks = optimize_chunk_size(page_text, 800)
                print(
                    f"  📄 Page {page_num + 1}: Generated {len(page_chunks)} chunks"
                )

                for chunk_idx, chunk_text in enumerate(page_chunks):
                    print(
                        f"  📄 Page {page_num + 1}, Chunk {chunk_idx + 1}: {len(chunk_text)} chars"
                    )
                    chunk_start = page_text.find(chunk_text)
                    if chunk_start != -1:
                        global_start = page_start + chunk_start
                        global_end = global_start + len(chunk_text)
                    else:
                        global_start = page_start
                        global_end = page_start + len(chunk_text)

                    chunks.append(
                        {
                            "text": chunk_text,
                            "page_number": page_num + 1,
                            "start_char": global_start,
                            "end_char": global_end,
                            "chunk_in_page": chunk_idx,
                        }
                    )
                    print(
                        f"  📄 Page {page_num + 1}, Chunk {chunk_idx + 1}: Added to chunks list"
                    )

            print(f"  📄 PDF processing completed, total chunks: {len(chunks)}")

            if len(chunks) == 0:
                print("  📄 No chunks generated from PDF pages, using fallback")
                simple_chunks = optimize_chunk_size(pos_text, 800)
                print(f"  📄 Fallback generated {len(simple_chunks)} chunks")
                for chunk in simple_chunks:
                    chunks.append(
                        {
                            "text": chunk,
                            "page_number": None,
                            "start_char": None,
                            "end_char": None,
                        }
                    )

        except Exception as e:
            print(f"  📄 PDF processing failed: {e}")
            print("  📄 Using fallback simple chunking")
            simple_chunks = optimize_chunk_size(text, 800)
            for chunk in simple_chunks:
                chunks.append(
                    {
                        "text": chunk,
                        "page_number": None,
                        "start_char": None,
                        "end_char": None,
                    }
                )
    else:
        # For non-PDF files, use page-aware chunking if explicit markers are present.
        marked_pages = parse_explicit_page_markers(text)
        if marked_pages:
            joined_text = "\n\n".join(str(p["text"]) for p in marked_pages)
            current_pos = 0
            for page in marked_pages:
                page_number = int(page["page_number"])  # type: ignore[arg-type]
                page_text = str(page["text"])
                page_start = joined_text.find(page_text, current_pos)
                if page_start == -1:
                    page_start = current_pos
                page_end = page_start + len(page_text)
                current_pos = page_end

                page_chunks = optimize_chunk_size(page_text, 800)
                for chunk_text in page_chunks:
                    chunk_start = page_text.find(chunk_text)
                    if chunk_start != -1:
                        global_start = page_start + chunk_start
                        global_end = global_start + len(chunk_text)
                    else:
                        global_start = page_start
                        global_end = page_start + len(chunk_text)
                    chunks.append(
                        {
                            "text": chunk_text,
                            "page_number": page_number,
                            "start_char": global_start,
                            "end_char": global_end,
                        }
                    )

            print(
                f"  📄 Parsed {len(marked_pages)} explicit pages from OCR markers; "
                f"generated {len(chunks)} chunks with page numbers"
            )
            print(f"  📄 Final result: returning {len(chunks)} chunks")
            return chunks

        # Fallback: no markers available, use simple chunking.
        if document_id:
            update_processing_status(
                document_id,
                "PROCESSING",
                "Chunking document text…",
                None,
                85,
            )
        simple_chunks = optimize_chunk_size(text, 800)
        for chunk in simple_chunks:
            chunks.append(
                {
                    "text": chunk,
                    "page_number": None,
                    "start_char": None,
                    "end_char": None,
                }
            )

    print(f"  📄 Final result: returning {len(chunks)} chunks")
    return chunks


def analyze_financial_data(df: pd.DataFrame) -> Dict[str, Any]:
    """Analyze financial data from DataFrame and return key metrics."""
    analysis = {}

    # Basic statistics for numeric columns
    numeric_cols = df.select_dtypes(include=["float64", "int64"]).columns
    if not numeric_cols.empty:
        analysis["numeric_summary"] = df[numeric_cols].describe().to_dict()

    # Try to identify financial metrics
    for col in df.columns:
        col_lower = col.lower()
        # Revenue/Income analysis
        if any(term in col_lower for term in ["revenue", "income", "sales"]):
            analysis["revenue_metrics"] = {
                "total": df[col].sum(),
                "average": df[col].mean(),
                "trend": df[col].pct_change().mean(),
            }
        # Expense analysis
        elif any(term in col_lower for term in ["expense", "cost", "spend"]):
            analysis["expense_metrics"] = {
                "total": df[col].sum(),
                "average": df[col].mean(),
                "trend": df[col].pct_change().mean(),
            }
        # Profit analysis
        elif any(term in col_lower for term in ["profit", "margin", "earnings"]):
            analysis["profit_metrics"] = {
                "total": df[col].sum(),
                "average": df[col].mean(),
                "trend": df[col].pct_change().mean(),
            }

    return analysis


def df_to_markdown(df: pd.DataFrame) -> str:
    """Convert a DataFrame to a markdown table string."""
    try:
        return df.to_markdown(index=False)
    except Exception:
        # Fallback to CSV if markdown fails
        return df.to_csv(index=False)


def clean_text_content(text: str) -> str:
    """Clean text while keeping paragraph breaks (old behavior collapsed all whitespace and hurt OCR/layout)."""
    if not text:
        return ""

    text = text.replace("\x00", "")
    text = re.sub(r"[\x01-\x08\x0B\x0C\x0E-\x1F\x7F]", "", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Collapse runs of spaces/tabs within a line; keep newlines between paragraphs
    lines = []
    for line in text.split("\n"):
        line = re.sub(r"[ \t]+", " ", line).strip()
        if line:
            lines.append(line)
    out = "\n\n".join(lines)
    # Trim excessive blank lines only
    out = re.sub(r"\n{3,}", "\n\n", out)
    return out.strip()


def process_file_background(
    file_content: bytes,
    file_extension: str,
    original_filename: str,
    userId: str,
    document_id: str,
):
    """Process file in background thread."""
    try:
        clear_cancel_processing(document_id)
        update_processing_status(
            document_id,
            "PROCESSING",
            f"Upload complete. Processing/training started for {original_filename}",
            None,
            0,
        )

        # Process the file with progress updates
        result = process_file_with_progress(
            file_content, file_extension, original_filename, userId, document_id
        )

        # Update status to completed
        update_processing_status(
            document_id,
            "COMPLETED",
            f"Successfully processed {original_filename}. Generated {result.get('chunks', 0)} chunks.",
            None,  # error
            100,  # progress
        )

    except ProcessingAborted as e:
        if e.reason == REASON_TIMEOUT:
            detail = (
                f"Processing exceeded {settings.DOCUMENT_PROCESSING_TIMEOUT_SEC}s. "
                "Try a smaller file, disable OCR_ALWAYS_FULLPAGE, or raise DOCUMENT_PROCESSING_TIMEOUT_SEC."
            )
        else:
            detail = "Processing was cancelled."
        update_processing_status(document_id, "CANCELLED", detail, e.reason, 0)
        log_to_backend("info", f"Processing stopped for {original_filename}: {e.reason}")

    except Exception as e:
        error_msg = f"Error processing {original_filename}: {str(e)}"
        update_processing_status(document_id, "ERROR", error_msg, str(e), 0)
        log_to_backend("error", error_msg, error=e)
    finally:
        clear_cancel_processing(document_id)


def process_file_with_progress(
    file_content: bytes,
    file_extension: str,
    original_filename: str,
    userId: str,
    document_id: str,
) -> dict:
    """Process file with progress tracking."""
    try:
        text_content = ""
        financial_analysis = {}

        timeout_sec = max(60, int(getattr(settings, "DOCUMENT_PROCESSING_TIMEOUT_SEC", 3600)))
        deadline = time.time() + timeout_sec
        abort_guard = make_abort_guard(document_id, deadline)

        update_processing_status(
            document_id, "PROCESSING", "Extracting text content...", None, 5
        )

        pdf_pages_for_chunking: Optional[List[str]] = None

        if file_extension.lower() == ".pdf":
            try:
                _peek = fitz.open(stream=file_content, filetype="pdf")
                total_pages = len(_peek)
                _peek.close()
            except Exception:
                total_pages = 0
            update_processing_status(
                document_id,
                "PROCESSING",
                f"Extracting text + OCR from {total_pages or '?'} PDF pages...",
                None,
                10,
            )

            # OCR is the single longest phase (8-12 min for big scanned PDFs).
            # Emit per-page progress so the UI bar moves continuously instead
            # of appearing frozen. OCR spans progress 10 -> 70 (60%
            # band), chunking 72-78, profile 80, embedding 82-98, upsert 99.
            _ocr_start_pct = 10
            _ocr_end_pct = 70
            _ocr_span = _ocr_end_pct - _ocr_start_pct
            _ocr_last_reported_pct: List[int] = [_ocr_start_pct]
            _ocr_last_reported_page: List[int] = [0]

            def _on_ocr_page(done: int, total: int) -> None:
                if total <= 0:
                    return
                pct = _ocr_start_pct + int((done / total) * _ocr_span)
                pct = max(_ocr_start_pct, min(pct, _ocr_end_pct))
                # Throttle: only emit when % changes OR every 5 pages,
                # whichever comes first. Avoids hammering the status lock
                # on tiny PDFs while still giving smooth feedback on big ones.
                should_report = (
                    pct != _ocr_last_reported_pct[0]
                    or (done - _ocr_last_reported_page[0]) >= 5
                    or done == total
                )
                if not should_report:
                    return
                _ocr_last_reported_pct[0] = pct
                _ocr_last_reported_page[0] = done
                update_processing_status(
                    document_id,
                    "PROCESSING",
                    f"OCR page {done}/{total}…",
                    None,
                    pct,
                )

            pdf_page_texts_raw = extract_pdf_page_texts(
                file_content,
                abort_guard=abort_guard,
                progress_callback=_on_ocr_page,
            )
            pdf_pages_for_chunking = [clean_text_content(p) for p in pdf_page_texts_raw]
            text_content = "\n\n".join(pdf_pages_for_chunking)
            log_to_backend(
                "info",
                f"PDF unified extract: {len(pdf_page_texts_raw)} pages, {len(text_content)} characters",
            )

        elif file_extension.lower() in [".docx", ".doc"]:
            update_processing_status(
                document_id, "PROCESSING", "Processing Word document...", None, 40
            )
            # Create a BytesIO object from the file content
            doc_stream = io.BytesIO(file_content)
            doc = docx.Document(doc_stream)
            for para in doc.paragraphs:
                text_content += para.text + "\n"

        elif file_extension.lower() == ".txt":
            update_processing_status(
                document_id, "PROCESSING", "Processing text file...", None, 40
            )
            # Handle plain text files
            try:
                # Try UTF-8 first
                text_content = file_content.decode("utf-8")
            except UnicodeDecodeError:
                try:
                    # Fallback to UTF-8 with error handling
                    text_content = file_content.decode("utf-8", errors="replace")
                except Exception:
                    # Final fallback to latin-1
                    text_content = file_content.decode("latin-1", errors="replace")

        elif file_extension.lower() in [".xlsx", ".xls", ".csv"]:
            update_processing_status(
                document_id, "PROCESSING", "Processing spreadsheet...", None, 40
            )
            file_stream = io.BytesIO(file_content)
            try:
                if file_extension.lower() == ".csv":
                    df = pd.read_csv(file_stream)
                    text_content = df_to_markdown(df)
                else:
                    xls = pd.ExcelFile(file_stream)
                    sheet_names = xls.sheet_names
                    sheet_tables = []
                    for idx, sheet in enumerate(sheet_names):
                        df_sheet = pd.read_excel(xls, sheet_name=sheet)
                        sheet_tables.append(
                            f"Sheet: {sheet}\n\n{df_to_markdown(df_sheet)}\n"
                        )
                        if idx == 0:
                            df = df_sheet  # Use first sheet for financial analysis
                    text_content = "\n\n".join(sheet_tables)
                # Perform financial analysis on the first sheet only
                financial_analysis = analyze_financial_data(df)
            except Exception as e:
                print(f"Error processing file: {str(e)}")
                raise ValueError(f"Error processing file: {str(e)}")

        elif file_extension.lower() in [
            ".png",
            ".jpg",
            ".jpeg",
            ".bmp",
            ".tiff",
            ".tif",
        ]:
            update_processing_status(
                document_id, "PROCESSING", "Processing image with OCR...", None, 40
            )
            # Handle image files with OCR
            if not OCR_ENABLED:
                raise ValueError(
                    f"OCR is disabled. Cannot process image file: {file_extension}"
                )

            try:
                image = Image.open(io.BytesIO(file_content))

                # Convert to RGB if necessary (for better OCR accuracy)
                if image.mode != "RGB":
                    image = image.convert("RGB")

                ocr_text = run_ocr_on_pil(image, lang=OCR_LANGUAGE)
                if ocr_text.strip():
                    text_content = ocr_text
                    log_to_backend(
                        "info",
                        f"OCR extracted {len(ocr_text)} characters from image {original_filename}",
                    )
                else:
                    log_to_backend(
                        "warning", f"No text found via OCR in image {original_filename}"
                    )
                    text_content = (
                        f"[Image file: {original_filename} - No text detected via OCR]"
                    )
            except Exception as e:
                log_to_backend(
                    "error",
                    f"Failed to process image file {original_filename}: {str(e)}",
                )
                raise ValueError(f"Error processing image file: {str(e)}")
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

        update_processing_status(
            document_id, "PROCESSING", "Cleaning and chunking text...", None, 72
        )

        # Clean text content before chunking
        cleaned_text_content = clean_text_content(text_content)

        # Debug: Check text content before chunking
        print(f"🔍 Text content before chunking:")
        print(f"  Content length: {len(cleaned_text_content)}")
        print(f"  Content preview: {cleaned_text_content[:200]}...")
        print(f"  Content is empty: {not cleaned_text_content.strip()}")

        # Use the same document ID that was passed from the upload endpoint
        document_id_for_storage = document_id

        abort_guard()
        # Optimize chunks with page tracking
        chunks_with_pages = optimize_chunk_size_with_pages(
            cleaned_text_content,
            file_extension,
            file_content,
            document_id,
            pdf_page_texts=pdf_pages_for_chunking,
        )
        total_chunks = len(chunks_with_pages)

        print(f"📊 Chunking results:")
        print(f"  Total chunks generated: {total_chunks}")
        if total_chunks > 0:
            for i, chunk in enumerate(chunks_with_pages[:3]):  # Show first 3 chunks
                print(
                    f"  Chunk {i + 1}: {len(chunk['text'])} chars - {chunk['text'][:50]}..."
                )
        else:
            print(f"  ⚠️ No chunks generated!")

        # ------------------------------------------------------------------
        # Document profile (generic metadata extraction).
        # Runs on the first ~5 pages of the already-OCR'd text. Used to seed
        # a header "summary" chunk, prepend an identity prefix on each body
        # chunk, and enrich every point's payload. See document_profile.py.
        # ------------------------------------------------------------------
        abort_guard()
        update_processing_status(
            document_id, "PROCESSING", "Extracting document profile…", None, 78
        )

        profile = empty_profile()
        try:
            if pdf_pages_for_chunking:
                profile_excerpt = "\n\n".join(pdf_pages_for_chunking[:5])
            else:
                profile_excerpt = cleaned_text_content[:20000]
            if profile_excerpt and profile_excerpt.strip():
                profile = extract_profile(profile_excerpt, original_filename)
                log_to_backend(
                    "info",
                    f"Document profile extracted: type='{profile.get('document_type')}', "
                    f"title='{profile.get('document_title')}', "
                    f"parties={len(profile.get('parties') or [])}, "
                    f"identifiers={len(profile.get('key_identifiers') or [])}",
                )
        except Exception as profile_err:
            log_to_backend(
                "warning",
                f"Document profile extraction failed (continuing without it): {profile_err}",
            )
            profile = empty_profile()

        chunk_prefix = render_chunk_prefix(profile, original_filename)
        profile_payload_fields = {
            "document_title": profile.get("document_title") or "",
            "document_type": profile.get("document_type") or "",
            "document_summary": profile.get("summary") or "",
            "parties": profile.get("parties") or [],
            "key_identifiers": profile.get("key_identifiers") or [],
            "key_topics": profile.get("key_topics") or [],
        }

        update_processing_status(
            document_id, "PROCESSING", "Storing in vector database...", None, 82
        )

        if qdrant_client:
            try:
                points = []

                # --- Summary / header chunk ------------------------------
                # Seeded so the query endpoint can always retrieve document
                # identity via a `chunk_kind == "summary"` filter, regardless
                # of how the user phrases a summarisation question.
                summary_text = render_summary_chunk_text(profile, original_filename)
                if summary_text and summary_text.strip():
                    try:
                        summary_embedding = get_cached_embedding(summary_text)
                        if hasattr(summary_embedding, "tolist"):
                            summary_embedding = summary_embedding.tolist()
                        points.append(
                            PointStruct(
                                id=str(uuid.uuid4()),
                                vector=summary_embedding,
                                payload={
                                    "document_id": document_id_for_storage,
                                    "user_id": userId,
                                    "content": summary_text,
                                    "filename": original_filename,
                                    "chunk_index": -1,
                                    "chunk_kind": "summary",
                                    "page_number": None,
                                    "start_char": None,
                                    "end_char": None,
                                    "file_type": file_extension.lower().lstrip("."),
                                    **profile_payload_fields,
                                },
                            )
                        )
                        print("  ✅ Summary chunk created")
                    except Exception as summary_err:
                        print(f"  ⚠️ Summary chunk creation failed: {summary_err}")

                # --- Body chunks -----------------------------------------
                # Embedding spans progress 82 -> 98 (16% band). Upsert is the
                # final step after this loop and flips the bar to 99 before
                # the outer completion handler flips it to 100.
                embed_total = len(chunks_with_pages)
                report_every = max(1, min(50, embed_total // 25 or 1))
                _embed_start_pct = 82
                _embed_end_pct = 98
                _embed_span = _embed_end_pct - _embed_start_pct
                for idx, chunk_data in enumerate(chunks_with_pages):
                    abort_guard()
                    if embed_total and (
                        idx == 0
                        or idx == embed_total - 1
                        or (idx + 1) % report_every == 0
                    ):
                        pct = _embed_start_pct + int(
                            ((idx + 1) / embed_total) * _embed_span
                        )
                        update_processing_status(
                            document_id,
                            "PROCESSING",
                            f"Embedding chunk {idx + 1}/{embed_total}…",
                            None,
                            min(pct, _embed_end_pct),
                        )
                    body_text = chunk_data["text"] or ""
                    # Inject a short identity header on each chunk so the
                    # retrieval surface always carries doc type + title,
                    # regardless of which chunk is returned by vector search.
                    stored_content = (
                        f"{chunk_prefix}\n{body_text}" if chunk_prefix else body_text
                    )
                    print(f"Processing chunk {idx + 1}/{embed_total}")
                    print(f"  Text length: {len(stored_content)}")

                    try:
                        embedding = get_cached_embedding(stored_content)
                        if hasattr(embedding, "tolist"):
                            embedding = embedding.tolist()
                    except Exception as embed_error:
                        print(f"  ❌ Embedding generation failed: {embed_error}")
                        continue

                    try:
                        point = PointStruct(
                            id=str(uuid.uuid4()),
                            vector=embedding,
                            payload={
                                "document_id": document_id_for_storage,
                                "user_id": userId,
                                "content": stored_content,
                                "filename": original_filename,
                                "chunk_index": idx,
                                "chunk_kind": "body",
                                "page_number": chunk_data.get("page_number"),
                                "start_char": chunk_data.get("start_char"),
                                "end_char": chunk_data.get("end_char"),
                                "file_type": file_extension.lower().lstrip("."),
                                **profile_payload_fields,
                            },
                        )
                        points.append(point)
                    except Exception as point_error:
                        print(f"  ❌ Point creation failed: {point_error}")
                        continue

                # Upsert with batched writes + retry.
                #
                # Why batched: a single bulk upsert of 500-700 points against
                # Qdrant Cloud has been observed to time out at the HTTP
                # layer. Splitting into ~100-point batches keeps each write
                # small and also lets us retry individual batches instead of
                # discarding the whole document on a single transient blip.
                #
                # Why fail loudly: before this change, a failed upsert was
                # caught and processing reported COMPLETED, leaving the DB
                # with a PROCESSED document that had zero vectors in Qdrant
                # — unqueryable but invisible to the user. Now a hard failure
                # propagates so ``process_file_background`` sets status=ERROR.
                if points:
                    update_processing_status(
                        document_id,
                        "PROCESSING",
                        f"Saving {len(points)} vectors to database…",
                        None,
                        99,
                    )
                    batch_size = max(
                        1,
                        int(
                            getattr(settings, "QDRANT_UPSERT_BATCH_SIZE", 100) or 100
                        ),
                    )
                    max_retries = max(
                        1,
                        int(
                            getattr(settings, "QDRANT_UPSERT_MAX_RETRIES", 3) or 3
                        ),
                    )
                    print(
                        f"Attempting to upsert {len(points)} points to Qdrant "
                        f"(batch_size={batch_size}, max_retries={max_retries})..."
                    )
                    total_points = len(points)
                    written = 0
                    for batch_start in range(0, total_points, batch_size):
                        abort_guard()
                        batch = points[batch_start : batch_start + batch_size]
                        batch_num = (batch_start // batch_size) + 1
                        total_batches = (
                            total_points + batch_size - 1
                        ) // batch_size
                        last_err: Optional[Exception] = None
                        for attempt in range(1, max_retries + 1):
                            try:
                                qdrant_client.upsert(
                                    collection_name=qdrant_collection,
                                    points=batch,
                                )
                                last_err = None
                                break
                            except Exception as upsert_err:
                                last_err = upsert_err
                                wait_s = min(10, 2 ** (attempt - 1))
                                print(
                                    f"  ⚠️ Batch {batch_num}/{total_batches} "
                                    f"attempt {attempt}/{max_retries} failed: "
                                    f"{upsert_err} (sleeping {wait_s}s)"
                                )
                                if attempt < max_retries:
                                    time.sleep(wait_s)
                        if last_err is not None:
                            raise RuntimeError(
                                f"Qdrant upsert failed after {max_retries} "
                                f"attempts on batch {batch_num}/{total_batches} "
                                f"({len(batch)} points): {last_err}"
                            ) from last_err
                        written += len(batch)
                        print(
                            f"  ✅ Batch {batch_num}/{total_batches} upserted "
                            f"({written}/{total_points} points)"
                        )
                    print(
                        f"✅ Successfully upserted {written} chunks "
                        f"(incl. header) for document {document_id_for_storage}"
                    )
                else:
                    print("⚠️ No valid points to upsert")

            except Exception as e:
                print(f"❌ Failed to upsert to Qdrant: {str(e)}")
                print(f"Error type: {type(e)}")
                import traceback

                print(f"Traceback: {traceback.format_exc()}")
                # Re-raise so the outer process_file_background handler marks
                # this document as ERROR instead of falsely reporting success.
                raise

        return {
            "chunks": total_chunks,
            "document_id": document_id_for_storage,
            "content": cleaned_text_content,
            "filename": original_filename,
            "file_size": len(file_content),
            "file_type": file_extension.lower().lstrip("."),
            "pages": len(
                [c for c in chunks_with_pages if c.get("page_number") is not None]
            ),
            "profile": profile,
        }

    except Exception as e:
        raise e


def process_file(
    file_content: bytes,
    file_extension: str,
    original_filename: str,
    userId: str,
) -> dict:
    """Process uploaded file with optimized chunking and Qdrant integration."""
    try:
        text_content = ""
        financial_analysis = {}
        pdf_pages_for_chunking: Optional[List[str]] = None

        if file_extension.lower() == ".pdf":
            try:
                _d = fitz.open(stream=file_content, filetype="pdf")
                _d.close()
            except fitz.FileDataError as e:
                log_to_backend(
                    "error", f"Failed to open PDF file '{original_filename}': {str(e)}"
                )
                raise ValueError(
                    f"PDF file '{original_filename}' appears to be corrupted or broken. Please try uploading a valid PDF file."
                )
            except Exception as e:
                log_to_backend(
                    "error",
                    f"Unexpected error opening PDF file '{original_filename}': {str(e)}",
                )
                raise ValueError(
                    f"Failed to process PDF file '{original_filename}': {str(e)}"
                )

            pdf_page_texts_raw = extract_pdf_page_texts(file_content)
            pdf_pages_for_chunking = [clean_text_content(p) for p in pdf_page_texts_raw]
            text_content = "\n\n".join(pdf_pages_for_chunking)
        elif file_extension.lower() in [".docx", ".doc"]:
            # Create a BytesIO object from the file content
            doc_stream = io.BytesIO(file_content)
            doc = docx.Document(doc_stream)
            for para in doc.paragraphs:
                text_content += para.text + "\n"
        elif file_extension.lower() == ".txt":
            # Handle plain text files
            try:
                # Try UTF-8 first
                text_content = file_content.decode("utf-8")
            except UnicodeDecodeError:
                try:
                    # Fallback to UTF-8 with error handling
                    text_content = file_content.decode("utf-8", errors="replace")
                except Exception:
                    # Final fallback to latin-1
                    text_content = file_content.decode("latin-1", errors="replace")
        elif file_extension.lower() in [".xlsx", ".xls", ".csv"]:
            file_stream = io.BytesIO(file_content)
            try:
                if file_extension.lower() == ".csv":
                    df = pd.read_csv(file_stream)
                    text_content = df_to_markdown(df)
                else:
                    xls = pd.ExcelFile(file_stream)
                    sheet_names = xls.sheet_names
                    sheet_tables = []
                    for idx, sheet in enumerate(sheet_names):
                        df_sheet = pd.read_excel(xls, sheet_name=sheet)
                        sheet_tables.append(
                            f"Sheet: {sheet}\n\n{df_to_markdown(df_sheet)}\n"
                        )
                        if idx == 0:
                            df = df_sheet  # Use first sheet for financial analysis
                    text_content = "\n\n".join(sheet_tables)
                # Perform financial analysis on the first sheet only
                financial_analysis = analyze_financial_data(df)
            except Exception as e:
                print(f"Error processing file: {str(e)}")
                raise ValueError(f"Error processing file: {str(e)}")
        elif file_extension.lower() in [
            ".png",
            ".jpg",
            ".jpeg",
            ".bmp",
            ".tiff",
            ".tif",
        ]:
            # Handle image files with OCR
            if not OCR_ENABLED:
                raise ValueError(
                    f"OCR is disabled. Cannot process image file: {file_extension}"
                )

            try:
                image = Image.open(io.BytesIO(file_content))

                # Convert to RGB if necessary (for better OCR accuracy)
                if image.mode != "RGB":
                    image = image.convert("RGB")

                ocr_text = run_ocr_on_pil(image, lang=OCR_LANGUAGE)
                if ocr_text.strip():
                    text_content = ocr_text
                    log_to_backend(
                        "info",
                        f"OCR extracted {len(ocr_text)} characters from image {original_filename}",
                    )
                else:
                    log_to_backend(
                        "warning", f"No text found via OCR in image {original_filename}"
                    )
                    text_content = (
                        f"[Image file: {original_filename} - No text detected via OCR]"
                    )
            except Exception as e:
                log_to_backend(
                    "error",
                    f"Failed to process image file {original_filename}: {str(e)}",
                )
                raise ValueError(f"Error processing image file: {str(e)}")
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

        # Clean text content before chunking
        cleaned_text_content = clean_text_content(text_content)

        # Debug: Check text content before chunking
        print(f"🔍 Text content before chunking:")
        print(f"  Content length: {len(cleaned_text_content)}")
        print(f"  Content preview: {cleaned_text_content[:200]}...")
        print(f"  Content is empty: {not cleaned_text_content.strip()}")

        # Generate a unique document ID for vector storage
        document_id = str(uuid.uuid4())

        # Optimize chunks with page tracking
        chunks_with_pages = optimize_chunk_size_with_pages(
            cleaned_text_content,
            file_extension,
            file_content,
            document_id,
            pdf_page_texts=pdf_pages_for_chunking,
        )
        total_chunks = len(chunks_with_pages)

        print(f"📊 Chunking results:")
        print(f"  Total chunks generated: {total_chunks}")
        if total_chunks > 0:
            for i, chunk in enumerate(chunks_with_pages[:3]):  # Show first 3 chunks
                print(
                    f"  Chunk {i + 1}: {len(chunk['text'])} chars - {chunk['text'][:50]}..."
                )
        else:
            print(f"  ⚠️ No chunks generated!")

        if qdrant_client:
            try:
                points = []
                for idx, chunk_data in enumerate(chunks_with_pages):
                    # Debug: Check chunk data
                    print(f"Processing chunk {idx + 1}/{len(chunks_with_pages)}")
                    print(f"  Text length: {len(chunk_data['text'])}")
                    print(f"  Text preview: {chunk_data['text'][:50]}...")

                    # Generate embedding with error handling
                    try:
                        embedding = get_cached_embedding(chunk_data["text"])
                        print(f"  ✅ Embedding generated: {len(embedding)} dimensions")
                        print(f"  Embedding type: {type(embedding)}")

                        # Ensure embedding is a list
                        if hasattr(embedding, "tolist"):
                            embedding = embedding.tolist()
                        print(f"  Final embedding type: {type(embedding)}")

                    except Exception as embed_error:
                        print(f"  ❌ Embedding generation failed: {embed_error}")
                        continue

                    # Create point with error handling
                    try:
                        point = PointStruct(
                            id=str(uuid.uuid4()),
                            vector=embedding,
                            payload={
                                "document_id": document_id,
                                "user_id": userId,
                                "content": chunk_data["text"],
                                "filename": original_filename,
                                "chunk_index": idx,
                                "page_number": chunk_data.get("page_number"),
                                "start_char": chunk_data.get("start_char"),
                                "end_char": chunk_data.get("end_char"),
                                "file_type": file_extension.lower().lstrip("."),
                            },
                        )
                        points.append(point)
                        print(f"  ✅ Point created successfully")

                    except Exception as point_error:
                        print(f"  ❌ Point creation failed: {point_error}")
                        continue

                # Upsert with error handling
                if points:
                    print(f"Attempting to upsert {len(points)} points to Qdrant...")
                    qdrant_client.upsert(
                        collection_name=qdrant_collection, points=points
                    )
                    print(
                        f"✅ Successfully upserted {len(points)} chunks to Qdrant for document {document_id}"
                    )
                else:
                    print("⚠️ No valid points to upsert")

            except Exception as e:
                print(f"❌ Failed to upsert to Qdrant: {str(e)}")
                print(f"Error type: {type(e)}")
                import traceback

                print(f"Traceback: {traceback.format_exc()}")
                # Continue processing even if Qdrant fails

        return {
            "chunks": total_chunks,
            "document_id": document_id,
            "content": cleaned_text_content,
            "filename": original_filename,
            "file_size": len(file_content),
            "file_type": file_extension.lower().lstrip("."),
            "pages": len(
                [c for c in chunks_with_pages if c.get("page_number") is not None]
            ),
        }

    except Exception as e:
        raise e


router = APIRouter()
security = HTTPBearer()


def derive_encryption_key(secret: str) -> bytes:
    """Derive the encryption key using HKDF, matching NextAuth.js implementation."""
    if not secret:
        raise ValueError("NEXTAUTH_SECRET is not set")

    # Convert the secret to bytes
    secret_bytes = secret.encode()

    # Use HKDF to derive the key
    hkdf = HKDF(
        algorithm=hashes.SHA256(),
        length=32,  # 32 bytes for AES-256
        salt=b"",  # NextAuth.js uses no salt
        info=b"NextAuth.js Generated Encryption Key",
        backend=default_backend(),
    )
    return hkdf.derive(secret_bytes)


def decrypt_token(token: str) -> dict:
    """Decrypt a NextAuth.js JWE token."""
    try:
        # Split the token into its components
        header_b64, _, iv_b64, ciphertext_b64, tag_b64 = token.split(".")

        # Decode the header
        header_padding = "=" * (-len(header_b64) % 4)
        header = json.loads(base64.urlsafe_b64decode(header_b64 + header_padding))

        if header.get("alg") != "dir" or header.get("enc") != "A256GCM":
            raise ValueError("Unsupported JWE algorithm or encryption")

        # Decode the other components
        iv = base64.urlsafe_b64decode(iv_b64 + "=" * (-len(iv_b64) % 4))
        ciphertext = base64.urlsafe_b64decode(
            ciphertext_b64 + "=" * (-len(ciphertext_b64) % 4)
        )
        tag = base64.urlsafe_b64decode(tag_b64 + "=" * (-len(tag_b64) % 4))

        # Derive the key using HKDF
        key = derive_encryption_key(settings.NEXTAUTH_SECRET)

        # Create AESGCM cipher
        aesgcm = AESGCM(key)

        # Decrypt the payload
        plaintext = aesgcm.decrypt(iv, ciphertext + tag, header_b64.encode())

        # Parse and return the decrypted payload
        return json.loads(plaintext)

    except Exception as e:
        print(f"Error decrypting token: {str(e)}")
        print(traceback.format_exc())
        raise


async def get_current_user(
    credentials: HTTPAuthorizationCredentials = Depends(security),
):
    """Get current user from Next.js session token."""
    try:
        token = credentials.credentials
        print(f"Received token: {token[:20]}...")

        # Enhanced debugging
        print(f"Token length: {len(token)}")
        print(f"Token parts: {len(token.split('.'))}")
        print(
            f"Token format: {'JWE' if len(token.split('.')) == 5 else 'JWT' if len(token.split('.')) == 3 else 'Unknown'}"
        )

        # Print NEXTAUTH_SECRET for debugging
        secret = settings.NEXTAUTH_SECRET
        print(f"Using NEXTAUTH_SECRET: {secret[:5] if secret else 'Not set'}...")
        print(f"NEXTAUTH_SECRET length: {len(secret) if secret else 0}")

        # Decrypt the token
        try:
            payload = decrypt_token(token)
            print("Successfully decrypted token")
            print(f"Decoded payload: {json.dumps(payload, indent=2)}")

            # Check if user is active
            user_id = payload.get("id")
            if user_id:
                # No database user lookup needed - AI service should not manage users
                pass

            return payload
        except Exception as e:
            print(f"Token decryption error: {str(e)}")
            raise HTTPException(status_code=401, detail=f"Invalid token: {str(e)}")

    except Exception as e:
        print(f"Authentication error: {str(e)}")
        print(traceback.format_exc())
        raise HTTPException(status_code=401, detail=f"Authentication failed: {str(e)}")


@router.post("/upload")
async def upload_files(
    files: Union[UploadFile, List[UploadFile]] = File(...),
    userId: Optional[str] = Form(None),
    description: Optional[str] = Form(None),
    category: Optional[str] = Form(None),
    api_key: str = Header(..., alias="X-API-Key"),
):
    """Upload one or more documents (PDF, DOC, DOCX, TXT, CSV, XLS, XLSX) for processing."""
    try:
        # Validate API key
        if api_key != settings.AI_SERVICE_API_KEY:
            raise HTTPException(status_code=401, detail="Invalid API key")

        # For now, we'll need to get user ID from the request or use a default
        # You might want to pass user ID in the form data or header
        if not userId:
            raise HTTPException(status_code=400, detail="User ID is required")

        log_to_backend("info", f"Processing upload for user: {userId}")

        # Ensure Qdrant collection exists before processing files (auto-create if allowed by API key)
        qsetup = ensure_qdrant_collection_exists()
        if not qsetup["ok"]:
            log_to_backend(
                "error",
                f"Qdrant setup failed: {qsetup.get('detail', 'unknown')}",
            )
            raise HTTPException(
                status_code=500,
                detail=qsetup.get("detail") or "Vector database setup failed",
            )

        # Handle both single file and list of files
        file_list = files if isinstance(files, list) else [files]
        uploaded_documents = []

        for file in file_list:
            log_to_backend("info", f"Starting upload for file: {file.filename}")

            # Read file content
            content = await file.read()

            # Debug logging
            log_to_backend(
                "info",
                f"File details: filename={file.filename}, size={file.size}, content_length={len(content) if content else 0}",
            )

            if not content:
                log_to_backend("error", f"File content is empty for {file.filename}")
                raise HTTPException(
                    status_code=400, detail=f"File '{file.filename}' content is empty"
                )

            # Validate file content
            if not content or len(content) == 0:
                raise HTTPException(
                    status_code=400,
                    detail=f"File '{file.filename}' is empty or could not be read",
                )

            # Check file size limit (50MB)
            if len(content) > 50 * 1024 * 1024:
                raise HTTPException(
                    status_code=400,
                    detail=f"File '{file.filename}' is too large. Maximum file size is 50MB.",
                )

            # Get file extension
            _, ext = os.path.splitext(file.filename)

            # Validate PDF files have proper header
            if ext.lower() == ".pdf":
                if not content.startswith(b"%PDF"):
                    raise HTTPException(
                        status_code=400,
                        detail=f"File '{file.filename}' does not appear to be a valid PDF file",
                    )
                # Check for minimum PDF size (PDF files should be at least 100 bytes)
                if len(content) < 100:
                    raise HTTPException(
                        status_code=400,
                        detail=f"File '{file.filename}' appears to be too small to be a valid PDF file",
                    )

            # Generate document ID for tracking
            document_id = str(uuid.uuid4())

            # Initialize processing status
            update_processing_status(
                document_id,
                "UPLOADED",
                f"File {file.filename} uploaded successfully",
                None,
                0,
            )

            # Start background processing
            processing_thread = threading.Thread(
                target=process_file_background,
                args=(content, ext, file.filename, userId, document_id),
            )
            processing_thread.daemon = True
            processing_thread.start()

            # Add to uploaded documents list
            uploaded_documents.append(
                {
                    "document_id": document_id,
                    "filename": file.filename,
                    "status": "PROCESSING",
                    "message": f"File uploaded successfully. Processing started in background.",
                }
            )

            log_to_backend(
                "info",
                f"Started background processing for {file.filename} with ID: {document_id}",
            )

        return {
            "message": f"Successfully uploaded {len(file_list)} files. Processing in background.",
            "documents": uploaded_documents,
            "processing_status": "BACKGROUND",
            "user": {
                "id": userId,
                "name": "User",
                "email": "user@example.com",
            },
        }
    except HTTPException:
        raise
    except Exception as e:
        log_to_backend("error", "Unexpected error in upload_files", error=e)
        raise HTTPException(
            status_code=500,
            detail="An unexpected error occurred while processing your upload",
        )


@router.post("/cancel/{document_id}")
async def cancel_document_processing(
    document_id: str,
    api_key: str = Header(..., alias="X-API-Key"),
):
    """Request cooperative cancellation of background processing (e.g. user deleted the document)."""
    if api_key != settings.AI_SERVICE_API_KEY:
        raise HTTPException(status_code=401, detail="Invalid API key")
    request_cancel_processing(document_id)
    log_to_backend(
        "info",
        f"Cancel requested for document processing job {document_id}",
    )
    return {
        "document_id": document_id,
        "message": "Cancellation requested; worker will stop at the next checkpoint.",
        "status": "CANCEL_REQUESTED",
    }


@router.post("/restart-processing/{document_id}")
async def restart_document_processing(
    document_id: str,
    api_key: str = Header(..., alias="X-API-Key"),
):
    """Restart processing for a document that appears to be stuck."""
    try:
        # Validate API key
        if api_key != settings.AI_SERVICE_API_KEY:
            raise HTTPException(status_code=401, detail="Invalid API key")

        # Check if document exists in processing status
        current_status = get_processing_status(document_id)
        if current_status["status"] == "NOT_FOUND":
            raise HTTPException(
                status_code=404, detail="Document not found in processing queue"
            )

        # Check if document is actually stuck (processing for too long)
        if current_status["status"] in ["COMPLETED", "ERROR"]:
            raise HTTPException(
                status_code=400, detail="Document is not stuck in processing"
            )

        # Refresh timestamp for monitors without clobbering live progress/details.
        # The background worker may still be running; resetting to 0 and "Restarting…"
        # made the UI look frozen. Preserve progress; normalize stale restart copy.
        raw_details = current_status.get("details")
        prev_details = raw_details.strip() if isinstance(raw_details, str) else ""
        if not prev_details or prev_details.startswith(
            "Restarting processing for document"
        ):
            prev_details = "Processing document…"

        prev_progress = current_status.get("progress")
        try:
            prev_p = int(prev_progress) if prev_progress is not None else 0
        except (TypeError, ValueError):
            prev_p = 0
        prev_p = max(0, min(99, prev_p))

        update_processing_status(
            document_id,
            "PROCESSING",
            prev_details,
            None,
            prev_p,
        )

        # Log the restart
        log_to_backend(
            "info",
            f"Document processing checkpoint acknowledged for {document_id} (monitor ping, progress preserved)",
        )

        return {
            "message": f"Processing checkpoint refreshed for document {document_id}",
            "document_id": document_id,
            "status": "RESTARTED",
            "timestamp": datetime.datetime.now().isoformat(),
        }

    except HTTPException:
        raise
    except Exception as e:
        log_to_backend(
            "error",
            f"Unexpected error restarting processing for {document_id}",
            error=e,
        )
        raise HTTPException(
            status_code=500,
            detail="An unexpected error occurred while restarting processing",
        )


@router.get("/status/{document_id}")
async def get_document_status(
    document_id: str,
    api_key: str = Header(..., alias="X-API-Key"),
):
    """Get processing status for a document."""
    try:
        # Validate API key
        if api_key != settings.AI_SERVICE_API_KEY:
            raise HTTPException(status_code=401, detail="Invalid API key")

        status = get_processing_status(document_id)
        # print(f"Status request for {document_id}: {status}")
        return {"document_id": document_id, "status": status}

    except HTTPException:
        raise
    except Exception as e:
        log_to_backend("error", "Unexpected error in get_document_status", error=e)
        raise HTTPException(
            status_code=500,
            detail="An unexpected error occurred while getting document status",
        )


@router.post("/embeddings/generate")
async def generate_embedding(
    text: str = Form(...),
    api_key: str = Header(..., alias="X-API-Key"),
):
    """Generate embedding for text using the same model as document processing."""
    try:
        # Validate API key
        if api_key != settings.AI_SERVICE_API_KEY:
            raise HTTPException(status_code=401, detail="Invalid API key")

        if not embedding_model:
            raise HTTPException(status_code=500, detail="Embedding model not available")

        # Generate embedding
        embedding = get_cached_embedding(text)

        return {
            "text": text,
            "embedding": embedding,
            "dimensions": len(embedding),
        }

    except HTTPException:
        raise
    except Exception as e:
        log_to_backend("error", "Unexpected error in generate_embedding", error=e)
        raise HTTPException(
            status_code=500,
            detail="An unexpected error occurred while generating embedding",
        )
